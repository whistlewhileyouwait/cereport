# 🎓 Final Clean Streamlit App for CEU Certificate + Email Message (PDF optional)

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pandas as pd
import os
from dotenv import load_dotenv
from supabase import create_client
from datetime import datetime
from docx2pdf import convert

from database import get_all_attendees, get_scan_log



# ─── Setup Supabase ─────────────────────────────────────
load_dotenv()
supabase = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_KEY"))


attendees_data = pd.DataFrame(get_all_attendees())
    if attendees_data.empty:
        st.error("⚠️ No attendees found in the database. Please check your Supabase connection.")
        st.stop()
scan_data = pd.DataFrame(get_scan_log())


# ─── Sessions Info ─────────────────────────────────────
sessions = [
    {"date": "May 2, 2025\n8:30", "title": "Prevention of Child Molestation", "speaker": "Matthew L. Ferrara, Ph.D.", "credits": 1.5},
    {"date": "May 2, 2025\n10:30", "title": "Overview of the Texas Department of Criminal Justice", "speaker": "Jennifer Deyne, LPC-S, LSOTP-S", "credits": 1.5},
    {"date": "May 2, 2025\n1:30 and 3:30", "title": "Taking the High Road-Ethical Challenges", "speaker": "Dan Powers, LCSW-S", "credits": 3.0},
    {"date": "May 3, 2025\n8:30", "title": "Role of LSOTPs & Polygraph Examiners", "speaker": "Sean Braun, LSOTP; Clay Wood", "credits": 1.5},
    {"date": "May 3, 2025\n10:30", "title": "Working with Female Juveniles", "speaker": "Francisco Torres, LSOTP", "credits": 1.5},
    {"date": "May 3, 2025\n1:30", "title": "Treating Clients with Mild Autism", "speaker": "Emily Dixon, Ph.D.", "credits": 1.5},
    {"date": "May 3, 2025\n3:30", "title": "Sexual Addiction Treatment", "speaker": "Matthew L. Ferrara, Ph.D.", "credits": 1.5},
    {"date": "May 4, 2025\n8:30", "title": "Risk Assessment Reports", "speaker": "Matthew L. Ferrara, Ph.D.", "credits": 1.5},
    {"date": "May 4, 2025\n10:30", "title": "Chaperon Program for Your Practice", "speaker": "Shelley Graham, Ph.D., LSOTP; Anna Shursen, Ph.D., LSOTP", "credits": 1.5},
    {"date": "May 4, 2025\n1:30", "title": "Legal Aspects of Deregistration", "speaker": "Scott Smith, Esq", "credits": 1.5},
    {"date": "May 4, 2025\n3:30", "title": "Adolescent Risk-Need-Responsivity", "speaker": "Casey O'Neal, Ph.D.", "credits": 1.5},
]

# ─── Streamlit UI ─────────────────────────────────────
st.title("🎓 CEU Certificate Generator")
badge_ids = sorted(attendees_data["badge_id"].astype(int).unique())
selected_badge = st.selectbox("Select attendee by badge number:", badge_ids)
    
selected_row = attendees_data[attendees_data["badge_id"] == selected_badge].iloc[0]
name = selected_row["name"]
email = selected_row.get("email", "no-email@example.com")

st.write(f"**Name:** {name}")
st.write(f"**Email:** {email}")


st.write("Attendees Columns:", attendees_data.columns.tolist())

def get_scans_by_day(scan_df, person_name):
    person_scans = scan_df[scan_df["name"] == person_name]
    times = person_scans["timestamp"].sort_values()

    summary = []
    for day, group in times.groupby(times.dt.date):
        check_in = group.min().strftime("%I:%M %p")
        check_out = group.max().strftime("%I:%M %p")
        summary.append((day.strftime("%B %d, %Y"), check_in, check_out))
    return summary

selected_sessions = []
for session in sessions:
    label = f"{session['date']} – {session['title']} ({session['credits']} hrs)"
    if st.checkbox(label):
        selected_sessions.append(session)



# ─── Generate Email Message ───────────────────────────
def generate_attendance_email(name, scans_df):
    person_scans = scans_df[scans_df["name"] == name]
    email_lines = [f"Hi {name.split()[0]},\n",
                   "Thank you so much for attending the conference! According to our scan records, it looks like you were present during the following times:\n"]
    for day in ["2025-05-02", "2025-05-03", "2025-05-04"]:
        day_dt = datetime.strptime(day, "%Y-%m-%d").date()
        day_scans = person_scans[person_scans["date"] == day_dt]
        if not day_scans.empty:
            in_time = day_scans["timestamp"].min().strftime("%I:%M %p")
            out_time = day_scans["timestamp"].max().strftime("%I:%M %p")
            email_lines.append(f"\u2022 {day_dt.strftime('%B %d, %Y')}: {in_time} to {out_time}")
        else:
            email_lines.append(f"\u2022 {day_dt.strftime('%B %d, %Y')}: No record")
    email_lines.append("\nIf any of these details need to be updated, just reply to this email and I’ll be happy to take care of it.\n")
    email_lines.append("Thanks again for being part of the event!\n\nBest,\n[Your Name]\n[Your Organization]")
    return "\n".join(email_lines)


# ─── Generate Certificate Function ────────────────────
def generate_certificate(name, sessions_attended, scans_df):
    doc = Document("Certficate of Training Blank (1).docx")
    total_credits = sum(s["credits"] for s in sessions_attended)
    for para in doc.paragraphs:
        if "Jane Doe" in para.text:
            para.text = para.text.replace("Jane Doe", name)
            run = para.runs[0] if para.runs else para.add_run(name)
            run.font.size = Pt(20)
            run.font.name = 'Monotype Corsiva'
            run.bold = True
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Monotype Corsiva')
        if "18 In-Person Hours" in para.text:
            para.text = f"{total_credits:.1f} In-Person Hours of Continuing Education Units"
            para.runs[0].font.size = Pt(12)

    table = doc.tables[0]
    for i in range(len(table.rows) - 1, 0, -1):
        table._tbl.remove(table._tbl.tr_lst[i])
    for s in sessions_attended:
        row = table.add_row()
        row.cells[0].text = s["date"]
        row.cells[1].text = s["title"]
        row.cells[2].text = s["speaker"]
        row.cells[3].text = f"{s['credits']} hrs."

    doc.add_page_break()
    doc.add_paragraph(generate_attendance_email(name, scans_df))

    docx_file = f"{name.replace(' ', '_')}_CERT.docx"
    doc.save(docx_file)
    convert(docx_file)  # Creates PDF automatically
    return docx_file, docx_file.replace(".docx", ".pdf")

# ─── Download Button ──────────────────────────────────
if name and selected_sessions:
    if st.button("🖨️ Generate Certificate with Email"):
        docx_file, pdf_file = generate_certificate(name, selected_sessions, scan_data)
        with open(pdf_file, "rb") as f:
            st.download_button("📥 Download PDF Certificate + Email", f, file_name=pdf_file)
        with open(docx_file, "rb") as f:
            st.download_button("📄 Download DOCX (editable)", f, file_name=docx_file)


